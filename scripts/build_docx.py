#!/usr/bin/env python3
"""Convert an Echo markdown draft into a Word document.

Supports the subset of markdown the template uses: #/##/### headings,
**bold** runs, pipe tables, bullet lists, and *italic* note lines.

Usage:
    python3 scripts/build_docx.py drafts/2026-06-12/echo_draft.md
Writes echo_draft.docx next to the input file.
"""
import pathlib
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_LINE_RE = re.compile(r"^\*(?!\*)(.+)\*$")


def add_runs(paragraph, text: str, italic: bool = False) -> None:
    """Split text on **bold** markers and add styled runs."""
    pos = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.italic = italic
        run = paragraph.add_run(match.group(1))
        run.bold = True
        run.italic = italic
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.italic = italic


def is_table_line(line: str) -> bool:
    return line.startswith("|") and line.rstrip().endswith("|")


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [split_row(l) for l in lines
            if not re.match(r"^\|[\s:|-]+\|$", l.strip())]
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c >= len(table.rows[r].cells):
                break
            cell = table.rows[r].cells[c]
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], cell_text)
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def build(md_path: pathlib.Path) -> pathlib.Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if is_table_line(line):
            block = []
            while i < len(lines) and is_table_line(lines[i].rstrip()):
                block.append(lines[i].rstrip())
                i += 1
            add_table(doc, block)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.lstrip().startswith(("- ", "* ")) and not ITALIC_LINE_RE.match(line.strip()):
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, line.lstrip()[2:])
        else:
            italic_match = ITALIC_LINE_RE.match(line.strip())
            para = doc.add_paragraph()
            if italic_match:
                add_runs(para, italic_match.group(1), italic=True)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                add_runs(para, line.strip())
        i += 1

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) != 2:
        sys.exit(__doc__)
    md = pathlib.Path(sys.argv[1])
    if not md.exists():
        sys.exit(f"Not found: {md}")
    print(f"Wrote {build(md)}")
